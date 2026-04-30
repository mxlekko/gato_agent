from __future__ import annotations

import base64
import mimetypes
import os
import platform
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


SUPPORTED_EXTENSIONS = {".docx", ".md", ".txt", ".pdf", ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}


@dataclass
class ParsedBlock:
    index: int
    text: str
    block_type: str
    char_count: int
    level: int | None = None


@dataclass
class ParsedDocument:
    doc_id: str
    file_name: str
    file_path: str
    source_type: str
    blocks: list[ParsedBlock]

    @property
    def text(self) -> str:
        return "\n\n".join(block.text for block in self.blocks)


def parse_document(file_path: str, doc_id: str) -> ParsedDocument:
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {suffix}")

    if suffix == ".docx":
        blocks = _read_docx_blocks(path)
    elif suffix == ".md":
        blocks = _read_markdown_blocks(path.read_text(encoding="utf-8", errors="ignore"))
    elif suffix == ".pdf":
        blocks = _read_pdf_blocks(path)
    elif suffix in _IMAGE_EXTENSIONS:
        blocks = _read_image_blocks(path)
    else:
        blocks = _read_text_blocks(path.read_text(encoding="utf-8", errors="ignore"))

    normalized_blocks = _normalize_blocks(blocks)
    if not normalized_blocks:
        raise ValueError(f"{path.name} does not contain readable text.")

    return ParsedDocument(
        doc_id=doc_id,
        file_name=path.name,
        file_path=str(path.resolve()),
        source_type=suffix.lstrip("."),
        blocks=normalized_blocks,
    )


def parse_text_content(
    text: str,
    doc_id: str,
    file_name: str,
    file_path: str,
    source_type: str,
    parse_mode: str = "markdown",
) -> ParsedDocument:
    if parse_mode == "markdown":
        blocks = _read_markdown_blocks(text)
    elif parse_mode == "text":
        blocks = _read_text_blocks(text)
    else:
        raise ValueError(f"Unsupported parse mode: {parse_mode}")

    normalized_blocks = _normalize_blocks(blocks)
    if not normalized_blocks:
        raise ValueError(f"{file_name} does not contain readable text.")

    return ParsedDocument(
        doc_id=doc_id,
        file_name=file_name,
        file_path=file_path,
        source_type=source_type,
        blocks=normalized_blocks,
    )


def blocks_to_editable_text(blocks: list[ParsedBlock]) -> str:
    rendered: list[str] = []
    for block in blocks:
        if block.block_type == "heading":
            level = min(max(block.level or 1, 1), 6)
            rendered.append(f"{'#' * level} {block.text}")
        elif block.block_type == "list":
            lines = [line.strip() for line in block.text.splitlines() if line.strip()]
            rendered.append("\n".join(_normalize_list_line(line) for line in lines))
        else:
            rendered.append(block.text)
    return "\n\n".join(rendered).strip()


def _read_markdown_blocks(text: str) -> list[ParsedBlock]:
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks: list[ParsedBlock] = []
    buffer: list[str] = []
    buffer_type = "paragraph"
    in_code_block = False

    def flush_buffer() -> None:
        nonlocal buffer, buffer_type
        if not buffer:
            return
        _append_block(blocks, "\n".join(buffer), buffer_type)
        buffer = []
        buffer_type = "paragraph"

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```") or stripped.startswith("~~~"):
            if in_code_block:
                buffer.append(line)
                flush_buffer()
                in_code_block = False
            else:
                flush_buffer()
                buffer_type = "code"
                buffer.append(line)
                in_code_block = True
            continue

        if in_code_block:
            buffer.append(line)
            continue

        if not stripped:
            flush_buffer()
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            flush_buffer()
            level = len(heading_match.group(1))
            _append_block(blocks, heading_match.group(2), "heading", level=level)
            continue

        if re.match(r"^(\s*[-*+]\s+|\s*\d+\.\s+)", stripped):
            if buffer and buffer_type != "list":
                flush_buffer()
            buffer_type = "list"
            buffer.append(stripped)
            continue

        if stripped.startswith("|") and stripped.count("|") >= 2:
            if buffer and buffer_type != "table":
                flush_buffer()
            buffer_type = "table"
            buffer.append(stripped)
            continue

        if buffer and buffer_type != "paragraph":
            flush_buffer()
        buffer_type = "paragraph"
        buffer.append(stripped)

    flush_buffer()
    return blocks


def _read_text_blocks(text: str) -> list[ParsedBlock]:
    segments = [segment.strip() for segment in re.split(r"\n\s*\n", text.replace("\r\n", "\n").replace("\r", "\n"))]
    blocks: list[ParsedBlock] = []
    for segment in segments:
        if segment:
            _append_block(blocks, segment, "paragraph")
    return blocks


def _read_docx_blocks(path: Path) -> list[ParsedBlock]:
    document = Document(str(path))
    blocks: list[ParsedBlock] = []

    for item in _iter_docx_items(document):
        if isinstance(item, Paragraph):
            text = item.text.strip()
            if not text:
                continue
            style_name = item.style.name if item.style is not None else ""
            heading_match = re.search(r"Heading\s+(\d+)", style_name, re.IGNORECASE)
            if heading_match:
                _append_block(blocks, text, "heading", level=int(heading_match.group(1)))
            elif style_name.lower().startswith("list"):
                _append_block(blocks, text, "list")
            else:
                _append_block(blocks, text, "paragraph")
        elif isinstance(item, Table):
            rows: list[str] = []
            for row in item.rows:
                cells = [" ".join(cell.text.split()) for cell in row.cells if cell.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                _append_block(blocks, "\n".join(rows), "table")

    return blocks


def _iter_docx_items(parent: DocxDocument | _Cell) -> Iterator[Paragraph | Table]:
    if isinstance(parent, DocxDocument):
        parent_element = parent.element.body
    else:
        parent_element = parent._tc

    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _normalize_blocks(blocks: list[ParsedBlock]) -> list[ParsedBlock]:
    normalized: list[ParsedBlock] = []
    for block in blocks:
        text = _normalize_text(block.text)
        if not text:
            continue
        normalized.append(
            ParsedBlock(
                index=len(normalized),
                text=text,
                block_type=block.block_type,
                char_count=len(text),
                level=block.level,
            )
        )
    return normalized


def _append_block(
    blocks: list[ParsedBlock],
    text: str,
    block_type: str,
    level: int | None = None,
) -> None:
    normalized = _normalize_text(text)
    if not normalized:
        return
    blocks.append(
        ParsedBlock(
            index=len(blocks),
            text=normalized,
            block_type=block_type,
            char_count=len(normalized),
            level=level,
        )
    )


def _normalize_text(text: str) -> str:
    lines = [line.rstrip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    cleaned: list[str] = []
    previous_blank = False
    for line in lines:
        stripped = " ".join(line.split())
        is_blank = not stripped
        if is_blank:
            if not previous_blank:
                cleaned.append("")
            previous_blank = True
            continue
        cleaned.append(stripped)
        previous_blank = False
    return "\n".join(cleaned).strip()


def _normalize_list_line(line: str) -> str:
    if re.match(r"^([-*+]\s+|\d+\.\s+)", line):
        return line
    return f"- {line}"


_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}


def _read_pdf_blocks(path: Path) -> list[ParsedBlock]:
    import fitz  # pymupdf

    blocks: list[ParsedBlock] = []
    try:
        doc = fitz.open(str(path))
    except Exception as exc:
        raise ValueError(f"Failed to open PDF {path.name}: {exc}") from exc

    try:
        for page_index in range(len(doc)):
            page = doc.load_page(page_index)
            page_blocks = _extract_pdf_text_blocks(page)
            if page_blocks and not _pdf_text_needs_ocr(page_blocks):
                _extend_blocks(blocks, page_blocks)
                continue

            page_text = _ocr_pdf_page(page, path.name, page_index)
            if page_text:
                _append_block(blocks, f"PDF 第 {page_index + 1} 页 OCR\n\n{page_text}", "paragraph")
            elif page_blocks:
                _extend_blocks(blocks, page_blocks)
    finally:
        doc.close()

    if not blocks:
        raise ValueError(f"{path.name}: 此 PDF 未检测到可选中的文字，OCR 也未识别到有效内容。")
    return blocks


def _extract_pdf_text_blocks(page: object) -> list[ParsedBlock]:
    blocks: list[ParsedBlock] = []
    text_blocks = page.get_text("blocks")
    for block in text_blocks:
        block_text = " ".join(block[4].split()) if len(block) >= 5 else str(block).strip()
        block_text = _clean_pdf_text(block_text)
        if block_text:
            _append_block(blocks, block_text, "paragraph")
    return blocks


def _extend_blocks(target: list[ParsedBlock], blocks: list[ParsedBlock]) -> None:
    for block in blocks:
        _append_block(target, block.text, block.block_type, level=block.level)


_PDF_FRAGMENTED_LATIN_RE = re.compile(r"(?<![A-Za-z])(?:[A-Za-z]\s+){2,}[A-Za-z](?![A-Za-z])")
_PDF_FRAGMENTED_DIGIT_RE = re.compile(r"(?<!\d)(?:\d\s+){2,}\d(?!\d)")
_PDF_LONG_NOISE_TOKEN_RE = re.compile(r"\b[A-Za-z0-9][A-Za-z0-9_-]{31,}\b")


def _pdf_text_needs_ocr(blocks: list[ParsedBlock]) -> bool:
    mode = os.environ.get("PDF_TEXT_EXTRACTION_MODE", "auto").strip().lower() or "auto"
    if mode in {"ocr", "force_ocr", "force-ocr"}:
        return True
    if mode in {"native", "text", "pymupdf"}:
        return False

    lines = [block.text for block in blocks if block.text.strip()]
    if not lines:
        return True

    suspicious_lines = 0
    fragmented_match_count = 0
    for line in lines:
        latin_matches = _PDF_FRAGMENTED_LATIN_RE.findall(line)
        digit_matches = _PDF_FRAGMENTED_DIGIT_RE.findall(line)
        line_score = len(latin_matches) + len(digit_matches)
        if line_score:
            suspicious_lines += 1
            fragmented_match_count += line_score

    line_ratio = suspicious_lines / max(len(lines), 1)
    return fragmented_match_count >= 5 or suspicious_lines >= 3 or line_ratio >= 0.18


def _clean_pdf_text(text: str) -> str:
    lines: list[str] = []
    for raw_line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        if _is_pdf_noise_line(line):
            continue
        line = _PDF_LONG_NOISE_TOKEN_RE.sub("", line)
        line = " ".join(line.split())
        if line:
            lines.append(line)
    return "\n".join(lines)


def _is_pdf_noise_line(line: str) -> bool:
    compact = re.sub(r"\s+", "", line)
    return len(compact) >= 32 and bool(re.fullmatch(r"[A-Za-z0-9_-]+", compact))


def _ocr_pdf_page(page: object, pdf_name: str, page_index: int) -> str:
    import fitz  # pymupdf

    render_scale = _read_float_env("PDF_OCR_RENDER_SCALE", default=2.5, minimum=1.0, maximum=4.0)
    pixmap = page.get_pixmap(matrix=fitz.Matrix(render_scale, render_scale), alpha=False)
    image_data = pixmap.tobytes("png")
    raw_image = _decode_pdf_ocr_image(image_data)
    preprocessed_image = _preprocess_pdf_ocr_image(raw_image)
    text = _run_pdf_ocr(raw_image, pdf_name=pdf_name, page_index=page_index, preprocessed_image=preprocessed_image)
    if text.strip() in {"空白页", "blank page", "Blank page"}:
        return ""
    return _clean_pdf_text(text)


def _read_int_env(name: str, default: int, minimum: int, maximum: int) -> int:
    try:
        parsed = int(os.environ.get(name, ""))
    except (TypeError, ValueError):
        return default
    if parsed < minimum:
        return minimum
    if parsed > maximum:
        return maximum
    return parsed


def _read_float_env(name: str, default: float, minimum: float, maximum: float) -> float:
    try:
        parsed = float(os.environ.get(name, ""))
    except (TypeError, ValueError):
        return default
    if parsed < minimum:
        return minimum
    if parsed > maximum:
        return maximum
    return parsed


def _decode_pdf_ocr_image(image_data: bytes) -> object:
    try:
        import cv2
        import numpy as np
    except ImportError as exc:
        raise ImportError(
            "OpenCV and numpy are required for scanned PDF OCR. "
            "Run `python -m pip install -r rag-service/requirements.txt`."
        ) from exc

    image_buffer = np.frombuffer(image_data, dtype=np.uint8)
    image = cv2.imdecode(image_buffer, cv2.IMREAD_COLOR)
    if image is None:
        raise ValueError("PDF 页面渲染图片解析失败，无法执行 OCR。")
    return image


def _preprocess_pdf_ocr_image(image: object) -> object:
    try:
        import cv2
        import numpy as np
    except ImportError as exc:
        raise ImportError(
            "OpenCV and numpy are required for scanned PDF OCR. "
            "Run `python -m pip install -r rag-service/requirements.txt`."
        ) from exc

    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    gray = _deskew_pdf_ocr_image(gray, cv2=cv2, np=np)
    denoised = cv2.fastNlMeansDenoising(gray, None, h=12, templateWindowSize=7, searchWindowSize=21)
    enhanced = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8)).apply(denoised)
    return cv2.adaptiveThreshold(
        enhanced,
        255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        31,
        11,
    )


def _deskew_pdf_ocr_image(gray_image: object, cv2: object, np: object) -> object:
    coords = np.column_stack(np.where(gray_image < 245))
    if len(coords) < 20:
        return gray_image

    angle = cv2.minAreaRect(coords)[-1]
    if angle < -45:
        angle = -(90 + angle)
    else:
        angle = -angle

    if abs(angle) < 0.2 or abs(angle) > 15:
        return gray_image

    height, width = gray_image.shape[:2]
    center = (width // 2, height // 2)
    matrix = cv2.getRotationMatrix2D(center, angle, 1.0)
    return cv2.warpAffine(
        gray_image,
        matrix,
        (width, height),
        flags=cv2.INTER_CUBIC,
        borderMode=cv2.BORDER_REPLICATE,
    )


def _run_pdf_ocr(image: object, pdf_name: str, page_index: int, preprocessed_image: object | None = None) -> str:
    engine = os.environ.get("PDF_OCR_ENGINE", "auto").strip().lower() or "auto"
    processed_image = preprocessed_image if preprocessed_image is not None else image
    if engine in {"macos", "ocrmac", "apple-vision", "apple_vision"}:
        text = _ocr_with_ocrmac(image)
        if text.strip() or preprocessed_image is None:
            return text
        return _ocr_with_ocrmac(preprocessed_image)
    if engine in {"paddle", "paddleocr"}:
        return _ocr_with_paddle(processed_image)
    if engine in {"tesseract", "pytesseract"}:
        return _ocr_with_tesseract(processed_image)
    if engine == "auto":
        errors: list[str] = []
        ocr_runners = [
            lambda: _ocr_with_tesseract(processed_image),
            lambda: _ocr_with_paddle(processed_image),
        ]
        if platform.system() == "Darwin":
            ocr_runners.insert(0, lambda: _ocr_with_ocrmac(image))
            if preprocessed_image is not None:
                ocr_runners.insert(1, lambda: _ocr_with_ocrmac(preprocessed_image))

        for ocr_runner in ocr_runners:
            try:
                text = ocr_runner()
            except Exception as exc:
                errors.append(str(exc))
                continue
            if text.strip():
                return text
        raise ValueError(f"{pdf_name} 第 {page_index + 1} 页本地 OCR 未识别到有效文本。{' | '.join(errors)}")

    raise ValueError("PDF_OCR_ENGINE 仅支持 auto、ocrmac、pytesseract 或 paddleocr。")


def _ocr_with_ocrmac(image: object) -> str:
    try:
        from ocrmac import ocrmac
    except ImportError as exc:
        raise ImportError(
            "ocrmac is required for macOS scanned PDF OCR. "
            "Run `python -m pip install -r rag-service/requirements.txt`."
        ) from exc

    image_path = _write_temp_ocr_image(image)
    try:
        language_preference = _ocrmac_language_preference()
        result = ocrmac.OCR(
            image_path,
            framework="vision",
            recognition_level=os.environ.get("PDF_OCR_RECOGNITION_LEVEL", "accurate").strip() or "accurate",
            language_preference=language_preference,
            confidence_threshold=_read_float_env("PDF_OCR_MIN_CONFIDENCE", default=0.0, minimum=0.0, maximum=1.0),
            detail=True,
        ).recognize()
    except Exception as exc:
        raise ValueError(f"macOS Vision OCR 识别失败: {exc}") from exc
    finally:
        Path(image_path).unlink(missing_ok=True)

    lines: list[str] = []
    for item in result:
        if not isinstance(item, (list, tuple)) or not item:
            continue
        text = str(item[0]).strip()
        if text:
            lines.append(text)
    return "\n".join(lines).strip()


def _ocrmac_language_preference() -> list[str] | None:
    raw_value = os.environ.get("PDF_OCR_LANG", "zh-Hans,en-US").strip()
    if not raw_value:
        return None
    return [item.strip() for item in raw_value.split(",") if item.strip()]


_PADDLE_OCR_READER: object | None = None


def _ocr_with_paddle(image: object) -> str:
    reader = _get_paddle_ocr_reader()
    min_confidence = _read_float_env("PDF_OCR_MIN_CONFIDENCE", default=0.3, minimum=0.0, maximum=1.0)
    image_path = _write_temp_ocr_image(image)

    try:
        try:
            result = reader.ocr(image_path, cls=True)
        except TypeError:
            result = reader.ocr(image_path)
    except Exception as exc:
        raise ValueError(f"PaddleOCR 识别失败: {exc}") from exc
    finally:
        Path(image_path).unlink(missing_ok=True)

    lines = _extract_paddle_ocr_lines(result, min_confidence=min_confidence)
    return "\n".join(lines).strip()


def _get_paddle_ocr_reader() -> object:
    global _PADDLE_OCR_READER
    if _PADDLE_OCR_READER is not None:
        return _PADDLE_OCR_READER

    _configure_paddle_cpu_runtime()
    try:
        from paddleocr import PaddleOCR
    except ImportError as exc:
        raise ImportError(
            "PaddleOCR is optional for scanned PDF OCR. Install paddleocr and paddlepaddle if PDF_OCR_ENGINE=paddleocr."
        ) from exc

    lang = os.environ.get("PDF_OCR_LANG", "ch").strip() or "ch"
    cpu_threads = _read_int_env("PDF_OCR_CPU_THREADS", default=1, minimum=1, maximum=8)
    try:
        _PADDLE_OCR_READER = PaddleOCR(
            use_angle_cls=True,
            lang=lang,
            show_log=False,
            use_gpu=False,
            enable_mkldnn=False,
            cpu_threads=cpu_threads,
        )
    except TypeError:
        _PADDLE_OCR_READER = PaddleOCR(use_angle_cls=True, lang=lang)
    return _PADDLE_OCR_READER


def _configure_paddle_cpu_runtime() -> None:
    os.environ.setdefault("FLAGS_use_mkldnn", "0")
    os.environ.setdefault("OMP_NUM_THREADS", "1")
    os.environ.setdefault("MKL_NUM_THREADS", "1")
    os.environ.setdefault("OPENBLAS_NUM_THREADS", "1")
    os.environ.setdefault("VECLIB_MAXIMUM_THREADS", "1")


def _write_temp_ocr_image(image: object) -> str:
    try:
        import cv2
    except ImportError as exc:
        raise ImportError("OpenCV is required for scanned PDF OCR.") from exc

    temp_file = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    temp_path = temp_file.name
    temp_file.close()
    if not cv2.imwrite(temp_path, image):
        raise ValueError("OCR 临时图片写入失败。")
    return temp_path


def _extract_paddle_ocr_lines(result: object, min_confidence: float) -> list[str]:
    lines: list[str] = []
    if not isinstance(result, list):
        return lines

    for page_result in result:
        if not page_result:
            continue
        for item in page_result:
            if not isinstance(item, (list, tuple)) or len(item) < 2:
                continue
            text_info = item[1]
            if not isinstance(text_info, (list, tuple)) or not text_info:
                continue
            text = str(text_info[0]).strip()
            confidence = 1.0
            if len(text_info) > 1:
                try:
                    confidence = float(text_info[1])
                except (TypeError, ValueError):
                    confidence = 1.0
            if text and confidence >= min_confidence:
                lines.append(text)

    return lines


def _ocr_with_tesseract(image: object) -> str:
    try:
        import pytesseract
        from PIL import Image as PILImage
    except ImportError as exc:
        raise ImportError(
            "pytesseract and Pillow are required for tesseract PDF OCR. "
            "Run `python -m pip install -r rag-service/requirements.txt`."
        ) from exc

    lang = os.environ.get("PDF_OCR_TESSERACT_LANG", "chi_sim+eng").strip() or "chi_sim+eng"
    config = os.environ.get("PDF_OCR_TESSERACT_CONFIG", "--psm 6").strip() or "--psm 6"
    try:
        return pytesseract.image_to_string(PILImage.fromarray(image), lang=lang, config=config).strip()
    except pytesseract.TesseractNotFoundError as exc:
        raise ValueError("未找到 tesseract 命令，请先安装 tesseract 和中文语言包。") from exc
    except pytesseract.TesseractError as exc:
        raise ValueError(f"Tesseract OCR 识别失败: {exc}") from exc


def _read_image_blocks(path: Path) -> list[ParsedBlock]:
    try:
        from PIL import Image as PILImage
    except ImportError:
        raise ImportError("Pillow is required for image parsing. Install it with: pip install Pillow")

    # Validate the image can be opened
    try:
        img = PILImage.open(str(path))
        img.verify()
    except Exception as exc:
        raise ValueError(f"{path.name} is not a readable image: {exc}") from exc

    description = _describe_image(path)
    if not description.strip():
        raise ValueError(f"{path.name}: 视觉模型未能生成有效描述。")

    blocks: list[ParsedBlock] = []
    _append_block(blocks, description, "paragraph")
    return blocks


def _describe_image(image_path: Path) -> str:
    with open(image_path, "rb") as fh:
        image_data = fh.read()

    mime_type = mimetypes.guess_type(str(image_path))[0] or "image/png"
    prompt = (
        "请详细描述这张图片的内容，包括所有文字信息、图表数据、对象、场景、人物、动作、颜色、布局等可见要素。"
        "如果图片包含文字，请完整转写。如果是图表，请描述数据趋势和关键数字。"
    )
    return _describe_image_data(
        image_data=image_data,
        mime_type=mime_type,
        prompt=prompt,
        max_tokens=_read_int_env("VISION_MAX_TOKENS", default=2048, minimum=512, maximum=8192),
    )


def _describe_image_data(image_data: bytes, mime_type: str, prompt: str, max_tokens: int) -> str:
    api_key = os.environ.get("DASHSCOPE_API_KEY", "").strip()
    if not api_key or api_key == "your_dashscope_api_key_here":
        raise ValueError("DashScope API Key 未配置，无法对图片进行视觉识别。请在 .env 中设置 DASHSCOPE_API_KEY。")

    vision_model = os.environ.get("VISION_MODEL", "qwen-vl-plus").strip()

    encoded_image = base64.b64encode(image_data).decode("ascii")
    data_uri = f"data:{mime_type};base64,{encoded_image}"

    try:
        from openai import OpenAI
    except ImportError:
        raise ImportError("openai is required for vision API calls.")

    client = OpenAI(
        api_key=api_key,
        base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
    )

    try:
        response = client.chat.completions.create(
            model=vision_model,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "image_url", "image_url": {"url": data_uri}},
                        {"type": "text", "text": prompt},
                    ],
                }
            ],
            max_tokens=max_tokens,
        )
    except Exception as exc:
        error_msg = str(exc)
        if "AccessDenied" in error_msg or "Unpurchased" in error_msg:
            raise ValueError(
                f"当前 API Key 没有 {vision_model} 模型的调用权限。请在百炼开通该模型。\n"
                f"原始错误: {error_msg}"
            ) from exc
        raise ValueError(f"视觉模型调用失败: {error_msg}") from exc

    content = response.choices[0].message.content
    return content.strip() if content else ""
